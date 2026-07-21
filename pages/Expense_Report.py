import base64
import hashlib
import json
from datetime import datetime
from io import BytesIO
from typing import Dict, List, Optional

import pandas as pd
import streamlit as st
from openai import OpenAI
from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="Expense Report Generator",
    page_icon="🧾",
    layout="wide",
)

st.title("🧾 Expense Report Generator")
st.caption(
    "Upload receipts • Extract expenses with AI • Review values • "
    "Exclude personal items • Download Word report"
)
st.divider()


# =========================================================
# OPENAI CLIENT
# =========================================================
def get_openai_client() -> Optional[OpenAI]:
    try:
        key = st.secrets.get("OPENAI_API_KEY", None)

        if not key:
            return None

        return OpenAI(api_key=key)

    except Exception:
        return None


client = get_openai_client()


# =========================================================
# CONSTANTS
# =========================================================
CATEGORY_OPTIONS = [
    "Trip Meals",
    "Airport Parking",
    "Green P Parking",
    "Company Supplies",
    "Hotel",
    "Fuel",
    "Tolls",
    "Transportation",
    "Other",
]

DEFAULT_COLUMNS = [
    "Receipt",
    "Date",
    "Merchant",
    "Category",
    "Receipt Total",
    "Excluded Amount",
    "Notes",
]


# =========================================================
# SESSION STATE
# =========================================================
if "expense_rows" not in st.session_state:
    st.session_state.expense_rows = []

if "processed_receipt_hashes" not in st.session_state:
    st.session_state.processed_receipt_hashes = {}

if "generated_expense_docx" not in st.session_state:
    st.session_state.generated_expense_docx = None

if "generated_expense_filename" not in st.session_state:
    st.session_state.generated_expense_filename = "expense_report.docx"


# =========================================================
# HELPERS
# =========================================================
def strip_text(value: Optional[str]) -> str:
    return str(value or "").strip()


def hash_bytes(content: bytes) -> str:
    if not content:
        return ""

    return hashlib.sha256(content).hexdigest()


def image_bytes_to_data_url(
    image_bytes: bytes,
    mime_type: str = "image/jpeg",
) -> str:
    encoded = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{encoded}"


def safe_float(value, default: float = 0.0) -> float:
    try:
        if value is None:
            return default

        if pd.isna(value):
            return default

        cleaned = str(value).replace("$", "").replace(",", "").strip()

        if cleaned == "":
            return default

        return round(float(cleaned), 2)

    except (TypeError, ValueError):
        return default


def parse_date(value) -> Optional[datetime]:
    if value is None:
        return None

    if isinstance(value, pd.Timestamp):
        return value.to_pydatetime()

    if isinstance(value, datetime):
        return value

    text = strip_text(value)

    if not text:
        return None

    formats = [
        "%Y-%m-%d",
        "%m/%d/%Y",
        "%d/%m/%Y",
        "%B %d, %Y",
        "%b %d, %Y",
    ]

    for fmt in formats:
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue

    try:
        parsed = pd.to_datetime(text, errors="coerce")

        if pd.isna(parsed):
            return None

        return parsed.to_pydatetime()

    except Exception:
        return None


def format_date_english(value) -> str:
    parsed = parse_date(value)

    if not parsed:
        return strip_text(value)

    return parsed.strftime("%B %d, %Y").replace(" 0", " ")


def clean_filename(text: str) -> str:
    allowed = []

    for char in text:
        if char.isalnum() or char in ("-", "_"):
            allowed.append(char)
        elif char == " ":
            allowed.append("_")

    cleaned = "".join(allowed).strip("_")
    return cleaned or "expense_report"


def normalize_category(category: str) -> str:
    category = strip_text(category)

    if category in CATEGORY_OPTIONS:
        return category

    lowered = category.lower()

    mapping = {
        "meal": "Trip Meals",
        "meals": "Trip Meals",
        "food": "Trip Meals",
        "restaurant": "Trip Meals",
        "airport": "Airport Parking",
        "airport parking": "Airport Parking",
        "green p": "Green P Parking",
        "green p parking": "Green P Parking",
        "parking": "Other",
        "supplies": "Company Supplies",
        "company material": "Company Supplies",
        "company supplies": "Company Supplies",
        "hotel": "Hotel",
        "fuel": "Fuel",
        "gas": "Fuel",
        "toll": "Tolls",
        "transportation": "Transportation",
    }

    return mapping.get(lowered, "Other")


# =========================================================
# AI RECEIPT EXTRACTION
# =========================================================
def extract_receipt_information(
    image_bytes: bytes,
    filename: str,
    mime_type: str,
) -> Dict:
    """
    Extract one expense from one receipt.

    The function does not automatically exclude personal items.
    It only lists possible personal/non-reimbursable items for review.
    """

    empty_result = {
        "date": "",
        "merchant": "",
        "category": "Other",
        "currency": "CAD",
        "receipt_total": 0.0,
        "possible_excluded_items": [],
        "notes": "",
    }

    if not client or not image_bytes:
        return empty_result

    data_url = image_bytes_to_data_url(
        image_bytes=image_bytes,
        mime_type=mime_type,
    )

    system_prompt = """
You extract expense information from Canadian receipts.

Return ONLY valid JSON using exactly this structure:

{
  "date": "YYYY-MM-DD",
  "merchant": "string",
  "category": "Trip Meals | Airport Parking | Green P Parking | Company Supplies | Hotel | Fuel | Tolls | Transportation | Other",
  "currency": "CAD",
  "receipt_total": 0.00,
  "possible_excluded_items": [
    {
      "item": "string",
      "amount_before_tax": 0.00,
      "estimated_amount_including_tax": 0.00,
      "reason": "string"
    }
  ],
  "notes": "string"
}

Rules:

1. Extract only information visible on the receipt.
2. Do not invent missing values.
3. receipt_total must be the final paid amount, including tax and tip.
4. Use CAD unless another currency is clearly shown.
5. Use Trip Meals for restaurants and travel food expenses.
6. Use Airport Parking only when the receipt clearly shows airport parking.
7. Use Green P Parking only when Green P is clearly shown.
8. Use Company Supplies for materials or supplies purchased for work.
9. Do not automatically remove any item from the total.
10. possible_excluded_items should only identify items that may be personal,
    such as soft drinks purchased alongside company materials.
11. If tax allocation for an item cannot be determined reliably,
    set estimated_amount_including_tax to 0.
12. If the receipt is unclear, explain the uncertainty in notes.
13. Date must be YYYY-MM-DD.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                f"Extract the expense information from this "
                                f"receipt. Filename: {filename}"
                            ),
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": data_url,
                                "detail": "high",
                            },
                        },
                    ],
                },
            ],
        )

        raw_content = response.choices[0].message.content or "{}"
        extracted = json.loads(raw_content)

        return {
            "date": strip_text(extracted.get("date")),
            "merchant": strip_text(extracted.get("merchant")),
            "category": normalize_category(
                extracted.get("category", "Other")
            ),
            "currency": strip_text(
                extracted.get("currency", "CAD")
            ).upper() or "CAD",
            "receipt_total": safe_float(
                extracted.get("receipt_total")
            ),
            "possible_excluded_items": extracted.get(
                "possible_excluded_items", []
            ) or [],
            "notes": strip_text(extracted.get("notes")),
        }

    except Exception as error:
        st.error(f"Could not extract {filename}: {error}")
        return empty_result


# =========================================================
# RECEIPT TO TABLE ROW
# =========================================================
def extraction_to_row(
    filename: str,
    extraction: Dict,
) -> Dict:
    possible_items = extraction.get("possible_excluded_items", [])

    possible_item_notes = []

    for item in possible_items:
        item_name = strip_text(item.get("item"))
        estimated_total = safe_float(
            item.get("estimated_amount_including_tax")
        )
        reason = strip_text(item.get("reason"))

        line = item_name

        if estimated_total > 0:
            line += f" — estimated CAD {estimated_total:.2f} incl. tax"

        if reason:
            line += f" ({reason})"

        if line:
            possible_item_notes.append(line)

    notes = strip_text(extraction.get("notes"))

    if possible_item_notes:
        possible_text = (
            "Possible excluded item: "
            + "; ".join(possible_item_notes)
        )

        notes = (
            f"{notes} {possible_text}".strip()
            if notes
            else possible_text
        )

    return {
        "Receipt": filename,
        "Date": extraction.get("date", ""),
        "Merchant": extraction.get("merchant", ""),
        "Category": extraction.get("category", "Other"),
        "Receipt Total": safe_float(
            extraction.get("receipt_total")
        ),
        "Excluded Amount": 0.0,
        "Notes": notes,
    }


# =========================================================
# REPORT DATA PROCESSING
# =========================================================
def prepare_expense_dataframe(
    edited_df: pd.DataFrame,
) -> pd.DataFrame:
    df = edited_df.copy()

    for column in DEFAULT_COLUMNS:
        if column not in df.columns:
            df[column] = ""

    df["Receipt Total"] = (
        pd.to_numeric(df["Receipt Total"], errors="coerce")
        .fillna(0)
        .round(2)
    )

    df["Excluded Amount"] = (
        pd.to_numeric(df["Excluded Amount"], errors="coerce")
        .fillna(0)
        .clip(lower=0)
        .round(2)
    )

    # Exclusion cannot be greater than receipt total.
    df["Excluded Amount"] = df[
        ["Receipt Total", "Excluded Amount"]
    ].min(axis=1)

    df["Reimbursable Amount"] = (
        df["Receipt Total"] - df["Excluded Amount"]
    ).round(2)

    df["Category"] = df["Category"].apply(normalize_category)

    df["Date Parsed"] = df["Date"].apply(parse_date)

    df = df[df["Reimbursable Amount"] > 0].copy()

    df = df.sort_values(
        by=["Date Parsed", "Receipt"],
        na_position="last",
    ).reset_index(drop=True)

    return df


def build_category_summary(
    df: pd.DataFrame,
) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(
            columns=["Category", "Total (CAD)"]
        )

    summary = (
        df.groupby("Category", as_index=False)["Reimbursable Amount"]
        .sum()
        .rename(
            columns={
                "Reimbursable Amount": "Total (CAD)"
            }
        )
    )

    summary["Total (CAD)"] = summary["Total (CAD)"].round(2)

    return summary


def determine_date_range(
    df: pd.DataFrame,
) -> tuple[str, str]:
    valid_dates = [
        value
        for value in df["Date Parsed"].tolist()
        if value is not None and not pd.isna(value)
    ]

    if not valid_dates:
        return "", ""

    first_date = min(valid_dates)
    last_date = max(valid_dates)

    return (
        format_date_english(first_date),
        format_date_english(last_date),
    )


# =========================================================
# WORD REPORT GENERATOR
# =========================================================
def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    align_right: bool = False,
):
    cell.text = ""

    paragraph = cell.paragraphs[0]

    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(10)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_header(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def generate_expense_word_report(
    recipient_name: str,
    df: pd.DataFrame,
    summary_df: pd.DataFrame,
    currency: str = "CAD",
    email_closing: str = "Thank you.",
    include_excluded_section: bool = True,
) -> bytes:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    first_date, last_date = determine_date_range(df)

    greeting = recipient_name.strip() or "Oren"

    paragraph = document.add_paragraph()
    paragraph.add_run(f"Hi {greeting},").bold = False

    document.add_paragraph()

    period_text = ""

    if first_date and last_date:
        period_text = (
            f" for the period from {first_date} to {last_date}"
        )

    intro = (
        "Please find attached the expense receipts"
        f"{period_text}. Below is a summary of the expenses:"
    )

    document.add_paragraph(intro)
    document.add_paragraph()

    # -----------------------------------------------------
    # DETAILED EXPENSE TABLE
    # -----------------------------------------------------
    expense_table = document.add_table(
        rows=1,
        cols=3,
    )

    expense_table.style = "Table Grid"
    expense_table.autofit = False

    expense_table.columns[0].width = Inches(2.0)
    expense_table.columns[1].width = Inches(3.2)
    expense_table.columns[2].width = Inches(1.4)

    headers = expense_table.rows[0].cells
    set_cell_text(headers[0], "Date", bold=True)
    set_cell_text(headers[1], "Category", bold=True)
    set_cell_text(
        headers[2],
        f"Amount ({currency})",
        bold=True,
        align_right=True,
    )

    for _, row in df.iterrows():
        cells = expense_table.add_row().cells

        set_cell_text(
            cells[0],
            format_date_english(row["Date"]),
        )

        set_cell_text(
            cells[1],
            strip_text(row["Category"]),
        )

        set_cell_text(
            cells[2],
            f'{safe_float(row["Reimbursable Amount"]):.2f}',
            align_right=True,
        )

    # -----------------------------------------------------
    # EXCLUDED PERSONAL EXPENSES
    # -----------------------------------------------------
    excluded_df = df[df["Excluded Amount"] > 0].copy()

    if include_excluded_section and not excluded_df.empty:
        document.add_paragraph()
        heading = document.add_paragraph()
        heading.add_run("Excluded Personal Expenses").bold = True

        excluded_table = document.add_table(
            rows=1,
            cols=3,
        )

        excluded_table.style = "Table Grid"

        excluded_headers = excluded_table.rows[0].cells

        set_cell_text(
            excluded_headers[0],
            "Date",
            bold=True,
        )
        set_cell_text(
            excluded_headers[1],
            "Description",
            bold=True,
        )
        set_cell_text(
            excluded_headers[2],
            f"Amount ({currency})",
            bold=True,
            align_right=True,
        )

        for _, row in excluded_df.iterrows():
            cells = excluded_table.add_row().cells

            set_cell_text(
                cells[0],
                format_date_english(row["Date"]),
            )

            description = strip_text(row["Notes"])

            if not description:
                description = "Personal expense, including applicable tax"

            set_cell_text(
                cells[1],
                description,
            )

            set_cell_text(
                cells[2],
                f'{safe_float(row["Excluded Amount"]):.2f}',
                align_right=True,
            )

    # -----------------------------------------------------
    # CATEGORY SUMMARY
    # -----------------------------------------------------
    document.add_paragraph()
    summary_heading = document.add_paragraph()
    summary_heading.add_run("Expense Summary").bold = True

    summary_table = document.add_table(
        rows=1,
        cols=2,
    )

    summary_table.style = "Table Grid"

    summary_headers = summary_table.rows[0].cells

    set_cell_text(
        summary_headers[0],
        "Category",
        bold=True,
    )

    set_cell_text(
        summary_headers[1],
        f"Total ({currency})",
        bold=True,
        align_right=True,
    )

    for _, row in summary_df.iterrows():
        cells = summary_table.add_row().cells

        set_cell_text(
            cells[0],
            row["Category"],
        )

        set_cell_text(
            cells[1],
            f'{safe_float(row["Total (CAD)"]):.2f}',
            align_right=True,
        )

    grand_total = round(
        df["Reimbursable Amount"].sum(),
        2,
    )

    total_cells = summary_table.add_row().cells

    set_cell_text(
        total_cells[0],
        "Grand Total",
        bold=True,
    )

    set_cell_text(
        total_cells[1],
        f"{grand_total:.2f}",
        bold=True,
        align_right=True,
    )

    document.add_paragraph()
    document.add_paragraph(email_closing.strip() or "Thank you.")

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output.getvalue()


# =========================================================
# SIDEBAR / REPORT SETTINGS
# =========================================================
with st.sidebar:
    st.header("Report Settings")

    recipient_name = st.text_input(
        "Recipient name",
        value="Oren",
    )

    currency = st.selectbox(
        "Currency",
        ["CAD", "USD"],
        index=0,
    )

    email_closing = st.text_input(
        "Email closing",
        value="Thank you.",
    )

    include_excluded_section = st.checkbox(
        "Show excluded personal expenses in Word",
        value=True,
    )

    st.info(
        "Review every extracted value before generating the report."
    )


# =========================================================
# RECEIPT UPLOAD
# =========================================================
st.subheader("1. Upload Receipts")

uploaded_receipts = st.file_uploader(
    "Upload receipt photos",
    type=["png", "jpg", "jpeg"],
    accept_multiple_files=True,
    key="expense_receipt_uploader",
)

if uploaded_receipts:
    st.caption(
        f"{len(uploaded_receipts)} receipt(s) selected."
    )

    preview_columns = st.columns(
        min(3, len(uploaded_receipts))
    )

    for index, uploaded_file in enumerate(uploaded_receipts):
        try:
            image = Image.open(
                BytesIO(uploaded_file.getvalue())
            )

            with preview_columns[index % len(preview_columns)]:
                st.image(
                    image,
                    caption=uploaded_file.name,
                    use_container_width=True,
                )

        except Exception:
            st.warning(
                f"Could not preview {uploaded_file.name}."
            )


# =========================================================
# ANALYZE RECEIPTS
# =========================================================
analyze_button = st.button(
    "🤖 Analyze Receipts",
    type="primary",
    disabled=not uploaded_receipts,
)

if analyze_button:
    if not client:
        st.error(
            "OPENAI_API_KEY is not configured in Streamlit secrets."
        )

    else:
        progress = st.progress(0)
        status = st.empty()

        total_files = len(uploaded_receipts)

        for index, uploaded_file in enumerate(uploaded_receipts):
            image_bytes = uploaded_file.getvalue()
            file_hash = hash_bytes(image_bytes)

            status.write(
                f"Analyzing {uploaded_file.name}..."
            )

            # Prevent duplicate extraction.
            if file_hash in st.session_state.processed_receipt_hashes:
                progress.progress((index + 1) / total_files)
                continue

            mime_type = uploaded_file.type or "image/jpeg"

            extraction = extract_receipt_information(
                image_bytes=image_bytes,
                filename=uploaded_file.name,
                mime_type=mime_type,
            )

            row = extraction_to_row(
                filename=uploaded_file.name,
                extraction=extraction,
            )

            st.session_state.expense_rows.append(row)

            st.session_state.processed_receipt_hashes[
                file_hash
            ] = uploaded_file.name

            progress.progress((index + 1) / total_files)

        status.empty()
        progress.empty()

        st.success("Receipt analysis completed.")
        st.rerun()


# =========================================================
# CLEAR DATA
# =========================================================
if st.session_state.expense_rows:
    if st.button("🗑️ Clear All Receipts"):
        st.session_state.expense_rows = []
        st.session_state.processed_receipt_hashes = {}
        st.session_state.generated_expense_docx = None
        st.rerun()


# =========================================================
# EDITABLE EXPENSE TABLE
# =========================================================
if st.session_state.expense_rows:
    st.divider()
    st.subheader("2. Review and Edit Expenses")

    st.caption(
        "Use Excluded Amount for personal items. "
        "Enter the amount including HST."
    )

    source_df = pd.DataFrame(
        st.session_state.expense_rows
    )

    for column in DEFAULT_COLUMNS:
        if column not in source_df.columns:
            source_df[column] = ""

    source_df = source_df[DEFAULT_COLUMNS]

    edited_df = st.data_editor(
        source_df,
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        column_config={
            "Receipt": st.column_config.TextColumn(
                "Receipt",
                disabled=True,
            ),
            "Date": st.column_config.TextColumn(
                "Date",
                help="Use YYYY-MM-DD",
            ),
            "Merchant": st.column_config.TextColumn(
                "Merchant",
            ),
            "Category": st.column_config.SelectboxColumn(
                "Category",
                options=CATEGORY_OPTIONS,
                required=True,
            ),
            "Receipt Total": st.column_config.NumberColumn(
                "Receipt Total",
                min_value=0.0,
                step=0.01,
                format="%.2f",
            ),
            "Excluded Amount": st.column_config.NumberColumn(
                "Excluded Amount incl. HST",
                min_value=0.0,
                step=0.01,
                format="%.2f",
                help=(
                    "Example: Diet Coke including its HST."
                ),
            ),
            "Notes": st.column_config.TextColumn(
                "Notes / excluded item",
                width="large",
            ),
        },
        key="expense_data_editor",
    )

    # Store edits.
    st.session_state.expense_rows = (
        edited_df.to_dict(orient="records")
    )

    final_df = prepare_expense_dataframe(edited_df)

    display_df = final_df[
        [
            "Date",
            "Merchant",
            "Category",
            "Receipt Total",
            "Excluded Amount",
            "Reimbursable Amount",
        ]
    ].copy()

    st.markdown("#### Reimbursable Values")

    st.dataframe(
        display_df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Receipt Total": st.column_config.NumberColumn(
                format="$ %.2f"
            ),
            "Excluded Amount": st.column_config.NumberColumn(
                format="$ %.2f"
            ),
            "Reimbursable Amount": st.column_config.NumberColumn(
                format="$ %.2f"
            ),
        },
    )

    # =====================================================
    # CATEGORY SUMMARY
    # =====================================================
    st.divider()
    st.subheader("3. Expense Summary")

    summary_df = build_category_summary(final_df)

    grand_total = round(
        final_df["Reimbursable Amount"].sum(),
        2,
    )

    summary_display = summary_df.copy()

    grand_total_row = pd.DataFrame(
        [
            {
                "Category": "Grand Total",
                "Total (CAD)": grand_total,
            }
        ]
    )

    summary_display = pd.concat(
        [summary_display, grand_total_row],
        ignore_index=True,
    )

    st.dataframe(
        summary_display,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Total (CAD)": st.column_config.NumberColumn(
                f"Total ({currency})",
                format="$ %.2f",
            ),
        },
    )

    metric_columns = st.columns(3)

    with metric_columns[0]:
        st.metric(
            "Receipts",
            len(final_df),
        )

    with metric_columns[1]:
        st.metric(
            "Excluded",
            f"{currency} {final_df['Excluded Amount'].sum():.2f}",
        )

    with metric_columns[2]:
        st.metric(
            "Grand Total",
            f"{currency} {grand_total:.2f}",
        )

    # =====================================================
    # REPORT PREVIEW
    # =====================================================
    st.divider()
    st.subheader("4. Email Preview")

    start_date, end_date = determine_date_range(final_df)

    if start_date and end_date:
        period_sentence = (
            f"for the period from {start_date} to {end_date}"
        )
    else:
        period_sentence = "for the selected period"

    preview_text = (
        f"Hi {recipient_name},\n\n"
        f"Please find attached the expense receipts "
        f"{period_sentence}. Below is a summary of the expenses:"
    )

    st.text_area(
        "Email introduction",
        value=preview_text,
        height=130,
        disabled=True,
    )

    # =====================================================
    # WORD GENERATION
    # =====================================================
    st.divider()
    st.subheader("5. Generate Word Report")

    if final_df.empty:
        st.warning(
            "Add at least one valid reimbursable expense."
        )

    else:
        generate_button = st.button(
            "📄 Generate Word Report",
            type="primary",
        )

        if generate_button:
            try:
                word_bytes = generate_expense_word_report(
                    recipient_name=recipient_name,
                    df=final_df,
                    summary_df=summary_df,
                    currency=currency,
                    email_closing=email_closing,
                    include_excluded_section=(
                        include_excluded_section
                    ),
                )

                first_date, last_date = determine_date_range(
                    final_df
                )

                filename_base = clean_filename(
                    f"Expense Report {first_date} to {last_date}"
                )

                st.session_state.generated_expense_docx = (
                    word_bytes
                )

                st.session_state.generated_expense_filename = (
                    f"{filename_base}.docx"
                )

                st.success(
                    "Word report generated successfully."
                )

            except Exception as error:
                st.error(
                    f"Could not generate Word report: {error}"
                )

        if st.session_state.generated_expense_docx:
            st.download_button(
                label="⬇️ Download Word Report",
                data=st.session_state.generated_expense_docx,
                file_name=(
                    st.session_state.generated_expense_filename
                ),
                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),
                type="primary",
            )

else:
    st.info(
        "Upload your receipts and click Analyze Receipts to begin."
    )